from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

import win32com.client as win32


ROOT = Path(r"C:\Users\Administrator\Desktop\wrj1")
OUTPUT_DIR = ROOT / "策划书"
DOCX_PATH = OUTPUT_DIR / "翼启全域-固定翼垂直起降无人机物流运输项目策划书.docx"
DOC_PATH = OUTPUT_DIR / "翼启全域-固定翼垂直起降无人机物流运输项目策划书.doc"


TITLE = "翼启全域"
SUBTITLE = "固定翼垂直起降无人机物流运输项目策划书"
COMPETITION = "全国大学生电子商务“创新、创意及创业”挑战赛参赛作品"
TEAM = "项目团队：翼启全域项目组"
DATE_TEXT = "编制时间：2026年3月"


ABSTRACT_PARAGRAPHS = [
    "翼启全域固定翼垂直起降无人机物流运输项目，是面向低空经济与现代物流融合发展趋势而提出的一项综合性创业方案。项目以固定翼垂直起降无人机为核心载体，通过集成机体结构、飞控导航、混合动力、模块化货舱、智能调度与安全冗余控制等关键能力，形成一套能够在山区、乡村、海岛、灾害现场和高时效配送场景中稳定运行的低空物流解决方案。",
    "当前我国物流体系在主干线运输方面已较为成熟，但在偏远地区末端配送、复杂地形跨越、应急物资快速投送以及高时效小批量运输方面，仍然存在覆盖不足、时效性差、综合成本高和突发情况下抗风险能力弱等问题。相较于传统多旋翼运载无人机航程短、固定翼飞行器依赖跑道的局限，本项目所采用的固定翼垂直起降方案兼具灵活部署和高效巡航两大优势，更适合承担真实物流任务。",
    "本项目的目标不是单纯制造一架能够飞行的无人机，而是构建“飞行平台 + 货运能力 + 航线运营 + 数据服务”一体化的低空运输体系。围绕乡村振兴、应急保障、海岛补给和城市高时效配送四大应用方向，项目计划在前期完成原型机研发、系统联调与试飞验证，中期推进示范航线建设和试点运营，后期形成多区域复制、平台化输出和生态协同发展。",
    "按照项目规划，设备平台将在最大航程、有效载重、起降便捷性和任务切换效率等方面形成明显竞争优势，并通过物流运营服务费、航线定制费、整机销售与租赁、平台授权及数据增值服务等方式形成多元化盈利来源。项目预计经过三年左右的市场培育和试点放量，具备形成区域性示范网络和稳定现金流的能力，兼具创新价值、创业潜力和较强的社会服务意义。",
]


SECTION_CONTENT = [
    {
        "h1": "一、计划书编制背景与目的",
        "items": [
            {
                "h2": "（一）低空物流产业发展背景",
                "paragraphs": [
                    "近年来，随着电子商务持续下沉、城乡流通体系不断完善以及应急保障体系建设逐步加强，社会对高效率、小批量、快响应物流服务的需求显著增加。传统地面运输虽然在大宗货运和常规城市配送中已形成较成熟体系，但在山区、海岛、边远乡村以及灾害中断场景下，仍然存在交通条件受限、运输路径冗长、时效无法保障和单位成本过高等问题。这些问题决定了未来物流体系必须向更加多层次、更加灵活和更加智能的方向演进。",
                    "与传统地面物流相比，低空物流能够跨越道路限制，压缩中间周转环节，在特定区域内形成点对点、站到站、急件优先的运输通道；与传统航空货运相比，低空物流具有起降灵活、部署周期短、基础设施要求低的特点；与普通多旋翼运载无人机相比，固定翼垂直起降无人机又兼具长航程、较高速和较强场景适应性的优势。因此，在当前低空经济进入应用验证期的背景下，发展固定翼垂直起降无人机物流运输项目具有较强的现实基础与产业意义。",
                ],
            },
            {
                "h2": "（二）计划书编制目的",
                "paragraphs": [
                    "本计划书的编制，首先是为了系统梳理翼启全域项目的市场机会、技术方案、商业模式、财务逻辑与实施路径，使项目从概念描述转向完整、清晰且具有执行依据的创业方案。通过对项目定位、目标客户、运营模式和风险控制等方面进行系统论证，可以为项目后续的研发推进、资源整合、校地合作、赛事路演与社会对接提供统一的表达框架。",
                    "其次，本计划书也是项目参加创新创业赛事、接受学校评审、争取合作伙伴支持和开展外部沟通的重要基础文件。对于赛事评委而言，计划书不仅需要说明项目“做什么”，更需要论证项目“为什么能做”“凭什么能做成”；对于潜在合作方而言，计划书要清楚呈现项目的目标市场、核心能力、实施条件和收益逻辑。因此，本计划书兼具展示、论证和落地三重功能。",
                ],
            },
            {
                "h2": "（三）项目基本概况",
                "paragraphs": [
                    "1．项目名称：翼启全域固定翼垂直起降无人机物流运输项目。",
                    "2．项目定位：面向偏远运输、应急投送与低空物流网络建设的智慧运输平台，重点服务乡村、山区、海岛及高时效配送场景。",
                    "3．项目宗旨：以技术创新提升物流效率，以场景落地创造社会价值，以平台化建设服务低空经济发展。",
                    "4．近期目标：完成原型机设计、飞控联调、样机试飞与典型场景验证，构建第一批示范性运输方案。",
                    "5．中长期目标：建立“整机装备 + 航线运营 + 平台服务”相结合的商业体系，形成可复制、可推广、可持续的低空物流运力网络。",
                ],
            },
        ],
    },
    {
        "h1": "二、市场定位",
        "items": [
            {
                "h2": "（一）低空物流市场的宏观分析",
                "paragraphs": [
                    "从宏观环境看，低空物流的发展同时受政策环境、市场需求、技术进步和社会治理方式变化四个维度驱动。一方面，低空经济正成为交通、制造、数字产业和区域经济融合发展的新增长点，地方政府对低空应用场景的关注度持续提升；另一方面，乡村振兴、区域协同发展、应急体系现代化和公共服务均等化，也在推动物流基础能力从传统地面网络向“地面 + 低空”融合网络延伸。",
                    "从需求结构看，当前最迫切的不是全行业、全品类的大规模替代，而是对高价值、高时效、难到达、强保障物流任务进行补位。偏远地区农产品上行和民生物资下行、灾后紧急救援、海岛和边远站点补给、医院样本和紧急药品运输等场景具有刚需、痛点明显、替代方案有限等特点，这为固定翼垂直起降无人机项目创造了较好的切入窗口。",
                ],
            },
            {
                "h2": "（二）目标市场的微观定位",
                "children": [
                    {
                        "h3": "1．产品定位",
                        "paragraphs": [
                            "本项目的产品定位不是一般意义上的展示型无人机，而是能够直接面向真实运输任务的低空物流平台。平台既包括固定翼垂直起降飞行器本体，也包括货运模块、任务调度、飞行监控、航线规划、运维保障和数据分析等配套能力，强调可执行、可运营和可扩展。",
                        ],
                    },
                    {
                        "h3": "2．客户定位",
                        "paragraphs": [
                            "项目的目标客户主要包括四类。第一类是政府和公共服务客户，如应急管理部门、乡镇政府、公共服务平台等，其核心需求是提升紧急运输能力和基层保障效率；第二类是农业与县域流通客户，如农业合作社、农产品基地、县域商贸中心等，其核心需求是提高农产品上行效率和缩短中间链路；第三类是医疗与专业配送客户，如医院、血站、疾控机构和检验机构等，其核心需求是对样本、药品和应急器械进行稳定、快速、定点运输；第四类是企业物流客户，如电商平台、园区企业和物流运营商等，其核心需求是提升复杂场景履约能力，降低部分高时效任务的综合成本。",
                        ],
                    },
                    {
                        "h3": "3．区域定位",
                        "paragraphs": [
                            "从区域角度看，项目优先布局地形复杂、道路条件不足、物流时效敏感度高且具有示范效应的区域，包括山地丘陵县域、岛屿及跨水域补给点、应急保障重点区域以及医疗高时效配送需求较高的城市节点。通过先进入典型区域，再复制推广至相似区域，可以有效降低前期市场进入成本并提高示范成功率。",
                        ],
                    },
                ],
            },
            {
                "h2": "（三）项目业务场景定位",
                "children": [
                    {
                        "h3": "1．乡村振兴与农产品上行",
                        "paragraphs": [
                            "在乡村物流场景中，固定翼垂直起降无人机可承担农产品样品运输、特色农产品快速外运、紧急生产资料投送和偏远乡村民生物资补给等任务。相较于地面运输，其优势在于能够减少弯绕路线、缩短运输时间并降低生鲜品类损耗，尤其适合山地和分散村落地区。",
                        ],
                    },
                    {
                        "h3": "2．应急物资运输",
                        "paragraphs": [
                            "在自然灾害、公共卫生事件、突发交通中断等场景下，项目可承担药品、食品、通信设备和应急工具的定点投送任务。由于无需依赖传统跑道，项目可在临时开阔地完成起降，从而在地面交通失效的情况下快速建立空中物资通道。",
                        ],
                    },
                    {
                        "h3": "3．海岛及偏远站点补给",
                        "paragraphs": [
                            "对于海岛、景区站点、边远值守点等区域，传统补给通常存在班次少、成本高、受天气影响大等问题。项目可通过定时巡航或按需起飞的方式，建立常态化补给能力，使这些区域具备更稳定的物流保障条件。",
                        ],
                    },
                    {
                        "h3": "4．城市高时效配送",
                        "paragraphs": [
                            "在城市应用中，本项目不以普遍替代传统快递为目标，而是优先聚焦医疗样本、紧急文件、高价值小件和特殊时限件的点对点运输。通过绕开地面拥堵、压缩任务路径和减少中间环节，项目可在高时效细分市场形成差异化服务能力。",
                        ],
                    },
                ],
            },
        ],
    },
    {
        "h1": "三、需求及竞争对手分析",
        "items": [
            {
                "h2": "（一）市场需求特征分析",
                "children": [
                    {
                        "h3": "1．服务对象具有明显的场景特殊性",
                        "paragraphs": [
                            "本项目所面向的客户并非传统大宗物流客户，而是存在运输难点和时间约束的特殊任务客户。这类客户往往更关注运输的稳定性、可到达性和任务响应速度，而不是单纯追求最低价格，因此更适合新型低空物流能力率先落地。",
                        ],
                    },
                    {
                        "h3": "2．业务需求具有强时效性和高可靠性要求",
                        "paragraphs": [
                            "无论是农产品样品快运、应急药品投送还是医疗标本运输，这类业务均要求较高的运输时效和履约稳定性。客户通常不能接受任务过程中频繁延误或中断，因此项目在飞控、冗余、安全和运营标准化方面必须具备较高要求，这也意味着项目一旦完成验证，将具备较强的客户粘性。",
                        ],
                    },
                    {
                        "h3": "3．现有替代方案无法充分满足需求",
                        "paragraphs": [
                            "传统地面物流依赖道路条件，遇到山区弯绕、道路中断和跨水域运输时效率显著下降；普通多旋翼运载无人机虽然部署灵活，但航程和载荷限制较明显，难以支撑中远距离持续运输任务；有人直升机虽然覆盖范围广，但运营成本高、调用难度大，不适合作为常态化低空物流工具。这意味着固定翼垂直起降无人机具备明显的市场补位空间。",
                        ],
                    },
                ],
            },
            {
                "h2": "（二）项目业务量预测",
                "paragraphs": [
                    "按照项目分阶段推进思路，前期以典型场景试点为主，中期形成区域化航线网络，后期逐步实现平台复制。结合项目定位和示范推进逻辑，对主要业务量做如下预测：",
                ],
                "table": {
                    "header": ["阶段", "重点区域/场景", "预计航线数量", "年任务架次", "业务特征"],
                    "rows": [
                        ["第1年", "样机验证与示范场景准备", "1-2条", "100-150架次", "以测试、验证、合作对接为主"],
                        ["第2年", "县域乡村和应急试点场景", "3-5条", "500-800架次", "初步形成运营收入"],
                        ["第3年", "区域化多场景联动", "8-12条", "2000-3000架次", "航线复制和客户结构扩展"],
                        ["第4-5年", "多区域网络化布局", "20条以上", "5000架次以上", "运营、设备与平台收入同步增长"],
                    ],
                },
            },
            {
                "h2": "（三）竞争对手分析",
                "children": [
                    {
                        "h3": "1．与传统地面物流企业相比",
                        "paragraphs": [
                            "传统物流企业在主干线路和成熟末端网络方面具有规模优势和组织优势，但在偏远地区、复杂地形和极端条件场景中，其履约效率受道路条件限制较大。翼启全域并不直接与传统物流网络全面竞争，而是选择其难以高效覆盖的任务场景进行切入，通过补位方式形成合作空间和差异化优势。",
                        ],
                    },
                    {
                        "h3": "2．与普通多旋翼运载无人机相比",
                        "paragraphs": [
                            "多旋翼无人机部署灵活、操控成熟，适合短距离、低载荷和精细化作业，但在航程、续航和中远距离运输效率方面存在明显不足。固定翼垂直起降无人机在保持灵活部署优势的同时，具备更长航程和更高巡航效率，更适用于区域间和中远距低空物流任务。",
                        ],
                    },
                    {
                        "h3": "3．与大型行业装备企业相比",
                        "paragraphs": [
                            "大型无人机制造企业和综合物流企业在资金、渠道和产业资源上具备明显优势，但往往更加重视标准化、大规模和成熟市场的应用布局。对于高校创新项目而言，重点不在于与其进行正面体量竞争，而在于以轻量、灵活和快速试点的方式切入细分场景，通过示范航线、定制化解决方案和区域深耕形成项目特色。",
                        ],
                    },
                ],
            },
            {
                "h2": "（四）项目核心竞争力",
                "paragraphs": [
                    "综合比较后，本项目的核心竞争力主要体现在以下几个方面：第一，技术结构上兼具垂直起降与固定翼巡航优势，兼顾部署灵活性和巡航效率；第二，项目面向的目标市场痛点明确，进入路径清晰，具有较强的场景针对性；第三，项目强调从设备能力向运营能力延伸，具有平台化发展思路；第四，项目兼具赛事展示性和实际应用性，便于争取政府、高校和企业多方合作资源；第五，项目具有较强社会价值，可与乡村振兴、公共服务和应急体系建设形成良好衔接。",
                ],
            },
        ],
    },
    {
        "h1": "四、企业发展战略",
        "items": [
            {
                "h2": "（一）战略愿景与发展目标",
                "paragraphs": [
                    "翼启全域的总体战略愿景，是建设具备示范效应和复制能力的低空物流运输平台，使固定翼垂直起降无人机真正进入区域流通、公共服务和应急保障体系之中。在短期内，项目以原型机研发、试飞验证和典型场景落地为重点；在中期内，项目以示范航线复制、运营标准建设和客户结构扩展为重点；在长期内，项目以形成“整机装备 + 运营服务 + 平台输出”协同发展的低空物流生态为目标。",
                ],
            },
            {
                "h2": "（二）战略重点",
                "children": [
                    {
                        "h3": "1．坚持技术自主与系统集成并重",
                        "paragraphs": [
                            "项目将围绕机体设计、飞控导航、动力系统、货舱模块和任务调度等核心能力持续优化，重点突破固定翼垂直起降过渡控制、复杂环境飞行稳定性和多场景任务适配等关键问题。通过持续研发投入和联合试验，不断提升平台成熟度和安全冗余水平。",
                        ],
                    },
                    {
                        "h3": "2．优先建设示范航线与样板场景",
                        "paragraphs": [
                            "在市场拓展上，项目不追求一开始大范围铺开，而是以“一个场景做深、一条航线做稳、一个区域做透”为原则，优先布局示范效果强、合作意愿高、任务痛点突出的区域。通过形成看得见、跑得通、能复用的样板方案，逐步带动周边区域复制。",
                        ],
                    },
                    {
                        "h3": "3．同步建立标准化运营体系",
                        "paragraphs": [
                            "低空物流要从试飞走向运营，必须建立面向实际业务的标准化运行机制。因此项目将同步推进飞前检查、任务审批、装载规范、飞行监控、异常处置、运维复盘和客户签收等制度建设，使技术能力最终转化为可靠的服务能力。",
                        ],
                    },
                    {
                        "h3": "4．构建合作生态与资源协同网络",
                        "paragraphs": [
                            "项目的发展不能仅靠单一团队独立完成，还需要地方政府、行业单位、科研团队、供应链企业和金融资源的协同参与。项目计划通过赛事路演、试点合作、校地协同和产业对接等方式，逐步构建“研发—验证—应用—推广”贯通的合作网络。",
                        ],
                    },
                ],
            },
            {
                "h2": "（三）阶段性实施战略",
                "table": {
                    "header": ["阶段", "时间安排", "战略重点", "主要成果"],
                    "rows": [
                        ["研发攻关期", "第1—6个月", "完成样机设计、飞控联调和核心系统验证", "形成原型机与关键模块"],
                        ["试飞验证期", "第7—12个月", "完成多场景试飞与性能优化", "形成可展示、可演示的样机体系"],
                        ["试点运营期", "第13—24个月", "建设示范航线并形成初步收入", "形成首批客户与运营数据"],
                        ["推广复制期", "第25—36个月", "复制航线、拓展区域、完善平台服务", "形成区域化网络和品牌影响力"],
                    ],
                },
            },
        ],
    },
    {
        "h1": "五、项目实施方案",
        "items": [
            {
                "h2": "（一）产品与技术方案",
                "children": [
                    {
                        "h3": "1．机体结构方案",
                        "paragraphs": [
                            "项目拟采用固定翼与垂直起降复合布局，机体以轻量化、高强度复合材料为主要结构基础，在满足载荷需求的同时兼顾气动效率和维护便利性。该方案能够有效支撑中远距巡航任务，并为不同场景货舱提供稳定安装基础。",
                        ],
                    },
                    {
                        "h3": "2．动力系统方案",
                        "paragraphs": [
                            "动力系统采用适合垂直起降与长航时巡航兼容的技术路线，使垂直起降阶段和巡航阶段分别获得最优能效支持。通过合理配置起降动力与巡航动力，项目能够兼顾起飞安全性、巡航经济性和整体续航能力，为区域低空运输提供更稳定的动力保障。",
                        ],
                    },
                    {
                        "h3": "3．飞控与导航方案",
                        "paragraphs": [
                            "飞控系统是本项目的核心技术单元之一。项目拟采用多传感器融合导航与自主飞控控制策略，将北斗/GPS、惯性导航、姿态感知和辅助视觉能力进行协同集成，提高路径保持精度、抗干扰能力和复杂环境适应性。围绕起降转换、巡航控制、任务返航和异常处置等关键环节建立完整控制逻辑，是项目技术实施的重点。",
                        ],
                    },
                    {
                        "h3": "4．货舱与任务模块方案",
                        "paragraphs": [
                            "为适配不同类型的运输任务，本项目采用模块化货舱设计思路，使标准物流箱、医药专用箱、保温冷链箱和应急物资箱能够实现快速拆装。通过统一接口标准和任务适配规范，可显著提升任务切换效率，减少人工装配时间并增强项目运营灵活性。",
                        ],
                    },
                    {
                        "h3": "5．安全冗余与避障方案",
                        "paragraphs": [
                            "安全始终是低空物流项目能否进入实际运营的前提条件。本项目在控制、感知、动力和操作流程上均强调冗余设计，拟通过飞行监测、异常告警、失联返航、航迹保护、环境感知和飞行前检查制度等方式构建多层次安全体系，确保项目既能飞起来，也能稳定、安全地飞下去。",
                        ],
                    },
                ],
            },
            {
                "h2": "（二）地面网络与节点建设",
                "paragraphs": [
                    "低空物流并不是脱离地面网络独立存在的系统，而是需要地面节点与空中节点协同运行。本项目计划建立“起降点—任务节点—保障节点”三类基础单元：起降点负责飞行器起降和任务准备；任务节点负责货物交接、装载和签收；保障节点负责补能、检修、监控和备份调度。前期可采用轻量化建设方式，在不增加过多固定资产投入的前提下满足试点运营需要，后期再根据业务规模逐步完善节点体系。",
                ],
            },
            {
                "h2": "（三）组织结构与运行机制",
                "paragraphs": [
                    "为了确保项目实施的系统性和执行力，项目拟设项目管理中心、技术研发中心、测试验证中心、市场运营中心和财务风控中心五个核心单元。项目管理中心负责总体统筹与对外协同；技术研发中心负责机体、飞控和软件系统开发；测试验证中心负责地面试验、飞行测试和数据记录；市场运营中心负责客户开发、试点落地和运营管理；财务风控中心负责预算、合规、融资及风险控制。通过明确职责边界和建立周例会、月复盘、节点验收等制度，提升团队协同效率。",
                ],
            },
            {
                "h2": "（四）业务启动与执行流程",
                "paragraphs": [
                    "项目业务启动流程按照“接单—评估—任务编排—装载—飞行—签收—复盘”的逻辑展开。首先由客户提出运输需求，项目团队根据货物类型、气象情况、航线条件和起降环境进行任务评估；评估通过后形成飞行任务单，完成货物装载和飞行前检查；任务执行过程中由地面站实时监控飞行状态，任务完成后由接收方签收并记录履约数据；最后由团队进行任务复盘，对飞行安全、时间消耗、货物状态和客户反馈进行记录，为后续运营优化提供依据。",
                ],
            },
        ],
    },
    {
        "h1": "六、商业模式",
        "items": [
            {
                "h2": "（一）品牌策略",
                "paragraphs": [
                    "“翼启全域”这一名称，既体现了固定翼垂直起降平台的技术属性，也体现了项目希望打通多场景、多区域运输能力的目标。品牌建设上，项目强调“专业、安全、稳定、可运营”四个关键词，不追求概念化表达，而是突出设备能力、任务执行和服务交付的真实可行性。通过统一品牌视觉、统一材料输出和统一案例表达，逐步形成项目在赛事、合作和试点中的辨识度。",
                ],
            },
            {
                "h2": "（二）产品与服务设计准则",
                "paragraphs": [
                    "本项目在产品与服务设计上坚持四项基本准则：一是以真实需求为导向，优先解决用户最痛的运输问题；二是以安全稳定为底线，不以牺牲安全换取表面效率；三是以标准化促进复制，尽量使设备接口、任务流程和服务规则统一；四是以可持续运营为目标，使项目从一开始就兼顾成本控制、收益逻辑和后续扩展空间。",
                ],
            },
            {
                "h2": "（三）经营模式",
                "children": [
                    {
                        "h3": "1．独立试点运营",
                        "paragraphs": [
                            "在项目早期，团队可围绕校地合作、乡村试点和典型示范任务开展独立试点运营，通过较小规模的任务执行积累飞行数据、客户反馈和运营经验。这种模式有利于项目快速验证技术与业务的匹配程度。",
                        ],
                    },
                    {
                        "h3": "2．合作共建模式",
                        "paragraphs": [
                            "当项目进入场景落地阶段后，可与地方政府、园区平台、物流企业和应急单位开展合作共建。合作方提供场景、节点和部分资源支持，项目方提供设备、运营和技术方案。该模式有助于降低前期市场进入成本，提高项目试点成功率。",
                        ],
                    },
                    {
                        "h3": "3．整机租赁与技术输出模式",
                        "paragraphs": [
                            "在项目进入推广阶段后，除了直接运营航线，还可向有明确需求的单位提供整机租赁、联合运营或技术输出服务。通过输出平台系统、调度能力和运营标准，项目可以将单一任务收入转化为更具持续性的技术服务收入。",
                        ],
                    },
                ],
            },
            {
                "h2": "（四）盈利模式",
                "paragraphs": [
                    "本项目的盈利来源主要由五部分构成。第一，物流运营服务费，即按照任务架次、航线距离、货物重量和保障等级向客户收取费用；第二，航线定制和试点方案设计费，即为地方政府、企业客户和公共服务单位提供专属运输方案并获得服务收入；第三，整机销售与租赁收入，即向合作单位提供设备交付和租赁服务；第四，平台系统授权收入，即输出飞控调度、任务管理和运营监控等软件能力；第五，数据增值服务收入，即围绕运维、航线优化和履约评价提供分析服务。多元化盈利结构有利于项目降低对单一收入来源的依赖，并提升中后期抗风险能力。",
                ],
            },
            {
                "h2": "（五）市场营销策略",
                "paragraphs": [
                    "市场营销方面，项目将坚持“案例先行、服务驱动、分层拓展”的策略。首先，通过高辨识度的试点案例建立项目品牌公信力，把样板航线和实飞成果作为最有说服力的推广材料；其次，通过赛事路演、成果展、媒体传播和学校平台扩大项目影响力，使更多潜在客户了解项目价值；再次，建立分层客户服务体系，对于政府和公共服务客户强调保障能力和社会价值，对于企业客户强调履约效率和运营收益，对于合作伙伴强调平台输出和联合发展的空间。项目价格策略也将采取先试点、后标准化的思路，在早期通过联合验证降低客户试用门槛，在中后期形成分层收费机制。",
                ],
            },
        ],
    },
    {
        "h1": "七、经济评价",
        "items": [
            {
                "h2": "（一）经济评价的主要指标",
                "paragraphs": [
                    "根据项目现阶段规划，翼启全域在启动期预计投入资金约800万元，重点用于样机研发、软件开发、测试验证、设备配置、试点运营及流动资金储备。项目计划在进入商业化试点后逐步形成收入，并在约18个月左右达到盈亏平衡。考虑到项目存在明显的技术研发周期和市场教育周期，前期财务目标以“验证商业可行性、建立可复制模型”为主，中后期再逐步释放规模效应。",
                ],
            },
            {
                "h2": "（二）经济评价的假设前提",
                "paragraphs": [
                    "本计划书中的财务数据，建立在以下假设基础之上：第一，项目能够按照既定时间节点完成样机研发和试飞验证，并进入典型场景试点；第二，试点区域能够提供相对稳定的任务需求和合作条件；第三，设备运维和人员成本在合理控制范围内；第四，项目采取边验证、边优化、边拓展的方式推进，不进行一次性大规模固定资产投入；第五，收入增长与航线数量扩展、客户结构丰富和平台化输出同步推进。",
                ],
            },
            {
                "h2": "（三）投资预算",
                "table": {
                    "header": ["项目", "金额（万元）", "用途说明"],
                    "rows": [
                        ["原型机研发与试制", "260", "机体结构设计、样机制造、部件集成"],
                        ["飞控算法与软件开发", "120", "飞控、任务调度、监控与数据平台开发"],
                        ["测试验证与试飞保障", "110", "测试场地、试飞组织、安全保障、耗材使用"],
                        ["设备与地面基础条件建设", "90", "地面站、工具设备、仓储及保障设施"],
                        ["市场试点与商务拓展", "80", "试点落地、商务对接、宣传展示"],
                        ["团队运营与管理费用", "90", "人员补贴、行政支出、日常管理"],
                        ["流动资金与风险储备", "50", "周转资金及不确定支出准备"],
                        ["合计", "800", "项目启动期总预算"],
                    ],
                },
            },
            {
                "h2": "（四）营业收入预测",
                "table": {
                    "header": ["年度", "预计营收（万元）", "预计成本（万元）", "预计利润（万元）", "说明"],
                    "rows": [
                        ["第1年", "200", "500", "-300", "研发、试飞、合作洽谈与示范准备阶段"],
                        ["第2年", "1500", "1200", "300", "试点航线启动并形成初步收入"],
                        ["第3年", "5000", "2800", "2200", "典型区域复制，客户结构扩展"],
                        ["第4年", "8000", "3500", "4500", "区域联动与平台输出并行增长"],
                        ["第5年", "15000", "5000", "10000", "形成网络效应与多元收入结构"],
                    ],
                },
            },
            {
                "h2": "（五）损益与回收分析",
                "paragraphs": [
                    "从测算结果看，项目在前期会经历较明显的研发投入期和市场培育期，因此第1年主要以投入为主，利润表现为负；随着第2年试点航线进入稳定运行，项目有望实现初步盈亏平衡；到第3年，随着业务量扩大、航线复制和服务结构丰富，项目的利润空间将快速释放。从长期看，只要项目能够持续提高设备利用率、控制运维成本并扩大客户复用率，整体盈利能力将不断增强。项目第五年营业收入目标达到1.5亿元，具备形成区域性低空物流品牌和稳定商业回报的可能性。",
                ],
            },
        ],
    },
    {
        "h1": "八、风险与控制",
        "items": [
            {
                "h2": "（一）政策与合规风险",
                "paragraphs": [
                    "低空物流项目受空域管理、飞行审批、运营资质和地方监管要求影响较大，一旦相关政策或审批机制发生变化，将对项目运行节奏和区域布局产生影响。对此，项目需要始终坚持合规先行原则，优先选择具备试点条件和政策支持基础的区域落地，同时提前准备相关资料，建立与地方主管部门的常态化沟通机制。",
                ],
            },
            {
                "h2": "（二）技术与安全风险",
                "paragraphs": [
                    "固定翼垂直起降无人机在系统集成、过渡控制、复杂环境飞行和长期可靠性方面均存在较高技术要求，若研发和验证不充分，可能影响项目后续运营安全。因此，项目必须坚持“先验证、后扩展”的思路，以多轮测试、分阶段放开和冗余设计为技术控制原则，把安全和可靠性放在商业化之前。",
                ],
            },
            {
                "h2": "（三）市场竞争风险",
                "paragraphs": [
                    "随着低空经济热度上升，传统物流企业、无人机设备企业和其他创业团队都有可能加速进入相关市场。如果项目在差异化定位、客户关系维护和标准化能力建设上推进不足，容易在后续竞争中被更大体量的参与者挤压。因此，项目必须坚持细分场景切入和样板航线先行策略，用场景深耕和运营能力形成壁垒。",
                ],
            },
            {
                "h2": "（四）资金与管理风险",
                "paragraphs": [
                    "创业项目普遍存在前期投入较大、回款周期较长和管理体系尚不完善的问题。如果在资金安排、团队分工、项目进度控制和合作管理方面缺乏系统性安排，就容易产生资源浪费或推进失序。对此，项目应实行分阶段预算管理和里程碑考核机制，严格控制固定支出，优先保障关键研发、试飞和试点所需资源。",
                ],
            },
            {
                "h2": "（五）风险控制措施",
                "paragraphs": [
                    "综合来看，本项目的风险控制应围绕“制度建设、技术验证、场景选择、资金节奏和合作机制”五个层面展开。制度建设方面，要形成飞行、运维、审批和应急处置规范；技术验证方面，要以多轮测试和迭代优化确保平台成熟度；场景选择方面，要优先进入需求明确、合作基础较好的试点区域；资金节奏方面，要分步融资、分步投放，避免超前扩张；合作机制方面，要通过校地合作、校企合作和联合运营降低项目独立承担的压力。只要项目能够在早期把握节奏、守住安全底线并建立标准化能力，整体风险将处于可控范围之内。",
                ],
            },
        ],
    },
]


def set_run_font(run, font_name: str, size: float, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def apply_normal_style(style):
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Pt(24)
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def apply_heading_style(style, font_name: str, size: float):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.bold = True
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_page_number(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_separate)
    paragraph._p.append(fld_end)


def set_table_cell_text(cell, text, font_name="宋体", size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold=False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_body_paragraph(document, text):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_heading(document, text, level):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = document.add_paragraph(style=style_map[level])
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", 12, True)
    elif level == 2:
        set_run_font(run, "黑体", 12, True)
    else:
        set_run_font(run, "黑体", 12, True)
    return p


def add_table(document, header, rows):
    table = document.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(header):
        set_table_cell_text(hdr_cells[idx], text, font_name="黑体", size=11)
    for row in rows:
        row_cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_table_cell_text(row_cells[idx], text, font_name="宋体", size=10.5)
    document.add_paragraph("")


def add_cover(document):
    for _ in range(6):
        document.add_paragraph("")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(COMPETITION)
    set_run_font(run, "黑体", 16, True)

    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(TITLE)
    set_run_font(run, "隶书", 28, True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(SUBTITLE)
    set_run_font(run, "隶书", 22, True)

    for _ in range(10):
        document.add_paragraph("")

    for text in (TEAM, DATE_TEXT):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        set_run_font(run, "隶书", 18, True)


def add_abstract(document):
    document.add_page_break()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("摘  要")
    set_run_font(run, "隶书", 22, True)

    for text in ABSTRACT_PARAGRAPHS:
        add_body_paragraph(document, text)

    p = document.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("关键词：固定翼垂直起降无人机；低空物流；应急运输；乡村物流；商业策划")
    set_run_font(run, "宋体", 12)


def add_toc_page(document):
    document.add_page_break()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("目  录")
    set_run_font(run, "隶书", 22, True)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("[[TOC]]")
    set_run_font(run, "宋体", 12)


def add_content(document):
    document.add_page_break()
    for section in SECTION_CONTENT:
        add_heading(document, section["h1"], 1)
        for item in section["items"]:
            add_heading(document, item["h2"], 2)
            for para in item.get("paragraphs", []):
                add_body_paragraph(document, para)
            if "table" in item:
                table_info = item["table"]
                add_table(document, table_info["header"], table_info["rows"])
            for child in item.get("children", []):
                add_heading(document, child["h3"], 3)
                for para in child.get("paragraphs", []):
                    add_body_paragraph(document, para)
                if "table" in child:
                    table_info = child["table"]
                    add_table(document, table_info["header"], table_info["rows"])


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.59)
    section.footer_distance = Inches(0.69)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    apply_normal_style(normal)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        apply_heading_style(doc.styles[name], "黑体", 12)

    if "TOC Style" not in [s.name for s in doc.styles]:
        toc_style = doc.styles.add_style("TOC Style", WD_STYLE_TYPE.PARAGRAPH)
        apply_normal_style(toc_style)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

    add_cover(doc)
    add_abstract(doc)
    add_toc_page(doc)
    add_content(doc)

    doc.save(DOCX_PATH)


def refresh_fields_and_export_doc():
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(str(DOCX_PATH))
    try:
        rng = doc.Content
        if rng.Find.Execute(FindText="[[TOC]]"):
            toc_range = rng.Duplicate
            toc_range.Text = ""
            doc.TablesOfContents.Add(
                Range=toc_range,
                UseHeadingStyles=True,
                UpperHeadingLevel=1,
                LowerHeadingLevel=3,
                UseFields=False,
                RightAlignPageNumbers=True,
                IncludePageNumbers=True,
                UseHyperlinks=True,
                HidePageNumbersInWeb=False,
            )
        if doc.TablesOfContents.Count >= 1:
            doc.TablesOfContents(1).Update()
        doc.Fields.Update()
        doc.Repaginate()
        doc.Save()
        doc.SaveAs(str(DOC_PATH), FileFormat=0)
    finally:
        doc.Close(False)
        word.Quit()


if __name__ == "__main__":
    build_document()
    refresh_fields_and_export_doc()
    print(str(DOCX_PATH))
    print(str(DOC_PATH))
